from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from reportlab.lib.pagesizes import A4
import os
from reportlab.lib.pagesizes import letter

# def converting_pdf(docx_file, pdf_file):
#     try:
#         if not docx_file.endswith(".docx"):
#             raise ValueError("File must be a .docx file")
            
#         doc = Document(docx_file)
#         pdf_doc = canvas.Canvas(pdf_file, pagesize=A4)

#         for block in doc.paragraphs:
#             if isinstance(block, Document.Paragraph):
#                 for run in block.runs:
#                     text = run.text
#                     if run.font.bold:
#                         pdf_doc.setFont('Times-Bold', 12)
#                     else:
#                         pdf_doc.setFont('Times-Roman', 12)
#                     pdf_doc.drawString(10*mm, 800-10*mm*len(text), text)

#             else:
#                 continue

#         pdf_doc.showPage()
#         pdf_doc.save()

#     except Exception as e:
#         print(f"Conversion failed: {e}")

def validate_file(file):
    """
    This function validates the file format.
    """
    if file.endswith(('.docx', '.doc', '.dotx', '.docm')):
        return True
    else:
        raise ValueError("Invalid file format. Please choose a .docx or .doc file.")

def converting_pdf(docx_file, pdf_file):
    """
    This function converts the docx file into a pdf file.
    """
    try:
        # Validate the file format
        validate_file(docx_file)

        # Open the docx file
        doc = Document(docx_file)

        # Create a PDF object
        pdf_gen = canvas.Canvas(pdf_file)

        # Get page size
        width, height = letter

        # Get the number of pages in the docx file
        num_pages = len(doc.element.part.get_pages())

        for i in range(num_pages):
            # Add text to the PDF
            for p in doc.paragraphs:
                pdf_gen.drawString(100, 750 - (i * 20), p.text)

            # Save the PDF page
            pdf_gen.showPage()

        # Close the PDF object
        pdf_gen.save()

    except Exception as e:
        print(f"Conversion failed: {e}")
        
        
doc1 = "C:/Users/pro/Desktop/test1.docx"
pdf1 = "C:/Users/pro/Desktop/test1.pdf"
test = converting_pdf(doc1, pdf1)
print(test)